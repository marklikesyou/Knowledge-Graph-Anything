import os
from dotenv import load_dotenv
from langchain.chat_models import ChatOpenAI
from langchain.graphs import Neo4jGraph
from langchain.graphs.graph_document import GraphDocument
from langchain.document_loaders import TextLoader
from langchain_experimental.graph_transformers.llm import LLMGraphTransformer
import io
import PyPDF2
import docx
from langchain.schema import Document

load_dotenv(".env.local")


class KnowledgeGraphAnything:
    def __init__(
        self,
        openai_api_key,
        neo4j_url,
        neo4j_username,
        neo4j_password,
        custom_schema=None,
    ):
        self.llm = ChatOpenAI(
            model="gpt-4o",
            openai_api_key=os.getenv("OPENAI_API_KEY"),
            temperature=0,
        )

        self.graph = Neo4jGraph(
            url=os.getenv("NEO4J_URL"),
            username=os.getenv("NEO4J_USERNAME"),
            password=os.getenv("NEO4J_PASSWORD"),
        )

        default_schema = """
        Extract meaningful relationships from the text while maintaining context:

        For each entity (node):
        1. Identify its main type/category
        2. Include relevant properties (name, date, description)
        3. Preserve any unique identifiers or key characteristics

        For each relationship:
        1. Specify the nature of the connection
        2. Include temporal context when available
        3. Add relevant details about the interaction
        4. Preserve any contextual information about impact or significance

        Focus on creating a comprehensive and interconnected knowledge graph while maintaining accuracy and context.
        """

        schema = custom_schema if custom_schema else default_schema

        self.transformer = LLMGraphTransformer(
            llm=self.llm,
            allowed_nodes=None,
            allowed_relationships=None,
        )

    def extract_text(self, file_content, file_type):
        try:
            try:
                text = file_content.decode("utf-8")
                return text
            except UnicodeDecodeError:
                if file_type == "pdf":
                    pdf_file = io.BytesIO(file_content)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text = " ".join(page.extract_text() for page in pdf_reader.pages)
                    return text
                elif file_type == "docx":
                    docx_file = io.BytesIO(file_content)
                    try:
                        doc = docx.Document(docx_file)
                        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                        return text
                    except Exception:
                        return None
                else:
                    return None
        except Exception:
            return None

    async def process_file_content(self, filename, content):
        file_type = filename.split(".")[-1].lower()
        text = self.extract_text(content, file_type)

        if not text:
            return []

        max_chunk_size = 12000
        chunks = [
            text[i : i + max_chunk_size] for i in range(0, len(text), max_chunk_size)
        ]

        all_graph_data = []
        for chunk in chunks:
            document = Document(page_content=chunk)
            try:
                graph_data = await self.transformer.aconvert_to_graph_documents(
                    [document]
                )
                if graph_data:
                    self.graph.add_graph_documents(graph_data, include_source=True)
                    all_graph_data.extend(graph_data)
            except Exception:
                pass

        return all_graph_data

    def clean_graph(self):
        try:
            query = "MATCH (n) DETACH DELETE n"
            self.graph.query(query)
        except Exception:
            pass

    async def process_files(self, uploaded_files):
        self.clean_graph()
        results = []
        for filename, content in uploaded_files.items():
            result = await self.process_file_content(filename, content)
            if result:
                results.append(result)
        return results

    def get_graph_statistics(self):
        try:
            node_count_query = "MATCH (n) RETURN count(n) as node_count"
            rel_count_query = "MATCH ()-[r]->() RETURN count(r) as rel_count"
            node_types_query = "MATCH (n) RETURN DISTINCT labels(n) as node_types"
            rel_types_query = "MATCH ()-[r]->() RETURN DISTINCT type(r) as rel_types"

            stats = {
                "nodes": self.graph.query(node_count_query)[0]["node_count"],
                "relationships": self.graph.query(rel_count_query)[0]["rel_count"],
                "node_types": [
                    r["node_types"] for r in self.graph.query(node_types_query)
                ],
                "relationship_types": [
                    r["rel_types"] for r in self.graph.query(rel_types_query)
                ],
            }
            return stats
        except Exception:
            return {
                "nodes": 0,
                "relationships": 0,
                "node_types": [],
                "relationship_types": [],
            }


def get_custom_schema():
    response = input(
        "Would you like to provide a custom schema for the knowledge graph? (yes/no)\n"
    ).lower()

    if response == "yes":
        print("\nEnter your custom schema below (press Enter twice when done):")
        lines = []
        while True:
            line = input()
            if line == "":
                break
            lines.append(line)

        return "\n".join(lines)

    return None


async def main():
    custom_schema = get_custom_schema()
    kg_builder = KnowledgeGraphAnything(
        openai_api_key=os.getenv("OPENAI_API_KEY"),
        neo4j_url=os.getenv("NEO4J_URL"),
        neo4j_username=os.getenv("NEO4J_USERNAME"),
        neo4j_password=os.getenv("NEO4J_PASSWORD"),
        custom_schema=custom_schema,
    )

    
    example_data_path = "./example data"
    if not os.path.exists(example_data_path):
        print(f"Error: {example_data_path} directory not found")
        return None, None

    uploaded_files = {}
    supported_extensions = {".txt", ".pdf", ".docx"}

    for filename in os.listdir(example_data_path):
        file_extension = os.path.splitext(filename)[1].lower()
        if file_extension in supported_extensions:
            file_path = os.path.join(example_data_path, filename)
            try:
                with open(file_path, "rb") as file:
                    uploaded_files[filename] = file.read()
                print(f"Loaded file: {filename}")
            except Exception as e:
                print(f"Error reading file {filename}: {str(e)}")
                continue

    if not uploaded_files:
        print(f"No supported files found in {example_data_path}")
        print(f"Supported file types: {', '.join(supported_extensions)}")
        return None, None

    print(f"\nProcessing {len(uploaded_files)} files...")
    results = await kg_builder.process_files(uploaded_files)
    stats = kg_builder.get_graph_statistics()

    print("\nKnowledge Graph Statistics:")
    print(f"Total Nodes: {stats['nodes']}")
    print(f"Total Relationships: {stats['relationships']}")
    print("\nNode Types:", stats["node_types"])
    print("Relationship Types:", stats["relationship_types"])

    print(
        "\nGraph creation complete! You can now view and query your knowledge graph in the Neo4j Browser."
    )
    return kg_builder, results


if __name__ == "__main__":
    import asyncio

    kg_builder, results = asyncio.run(main())
